#!/usr/bin/env python3
"""
Script to convert text resumes to DOCX format
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import glob

def create_docx_from_text(text_file):
    """Convert a text file to DOCX format"""
    docx_file = text_file.replace('.txt', '.docx')
    
    # Read text file
    with open(text_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process each line
    lines = content.split('\n')
    for i, line in enumerate(lines):
        line = line.strip()
        
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue
        
        # Check if it's a header (all caps or title case, short line)
        is_header = (
            line.isupper() or 
            (len(line) < 50 and line[0].isupper() and '|' not in line) or
            line.startswith('EXPERIENCE') or
            line.startswith('EDUCATION') or
            line.startswith('SKILLS') or
            line.startswith('EXPERIENCIA') or
            line.startswith('EDUCACIÓN') or
            line.startswith('HABILIDADES') or
            line.startswith('WORK EXPERIENCE') or
            line.startswith('CAREER HISTORY') or
            line.startswith('PROFESSIONAL SUMMARY') or
            line.startswith('TECHNICAL SKILLS') or
            line.startswith('CERTIFICATIONS') or
            line.startswith('PROJECTS') or
            line.startswith('PUBLICATIONS') or
            line.startswith('AWARDS') or
            line.startswith('LANGUAGES')
        )
        
        # Check if it's a name/contact line (first few lines, contains email/phone)
        is_contact = (
            i < 3 and 
            ('@' in line or 'phone' in line.lower() or 'tel' in line.lower() or 
             'linkedin' in line.lower() or 'location' in line.lower() or 'ubicación' in line.lower())
        )
        
        # Check if it's a job title line (contains | separator)
        is_job_title = '|' in line and len(line.split('|')) >= 2
        
        # Check if it's a bullet point
        is_bullet = line.startswith('•') or line.startswith('-') or line.startswith('*')
        
        # Add paragraph with appropriate formatting
        p = doc.add_paragraph()
        
        if is_contact or (i == 0 and len(line) < 100):
            # Name or contact info - center and bold
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14) if i == 0 else Pt(11)
        elif is_header:
            # Section header - bold, larger
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            p.space_after = Pt(6)
        elif is_job_title:
            # Job title line - bold
            run = p.add_run(line)
            run.bold = True
            p.space_after = Pt(3)
        elif is_bullet:
            # Bullet point
            p.style = 'List Bullet'
            p.add_run(line.lstrip('•-* '))
            p.space_after = Pt(3)
        else:
            # Regular text
            p.add_run(line)
            p.space_after = Pt(3)
    
    # Save document
    doc.save(docx_file)
    print(f"Created: {docx_file}")
    return docx_file

def main():
    # Get all text files in current directory
    text_files = sorted(glob.glob('resume-*.txt'))
    
    if not text_files:
        print("No resume text files found!")
        return
    
    print(f"Found {len(text_files)} text files to convert...")
    
    for text_file in text_files:
        try:
            create_docx_from_text(text_file)
        except Exception as e:
            print(f"Error converting {text_file}: {e}")

if __name__ == '__main__':
    main()














